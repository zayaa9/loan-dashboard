# -*- coding: utf-8 -*-
"""
report_cowork.py — ETL-ийн эцсийн датанд үндэслэн ЦОГЦ Word тайлан үүсгэнэ.
=====================================================================
Бүтэц (тоонууд долоо хоног бүр датанаас автоматаар шинэчлэгдэнэ):
  Тойм → 1.Онооны тархалт → 2.Хугацаа хэтрэлт (онооны бүс 30 оноогоор, 1+/15+ хэтрэлт)
  → 3.Хэтэрсэн vs хэтрээгүй → 4.Нас×хүйсний 6 сегмент → Дүгнэлт
  → Хавсралт: сегмент × хэтрэлтийн бүсийн хамаарлын хүснэгт (top 5)

ЧУХАЛ: archive/*.parquet (etl_cowork.py-ийн эцсийн дата) дээр ажиллана.
       run_loan_etl.bat дотор etl_cowork.py-ийн ДАРАА ажиллана.

Шаардлагатай сан:
    pip install pandas matplotlib python-docx pyarrow

Ажиллуулах:
    py report_cowork.py
"""

import os
import io
import glob
import datetime as dt
from pathlib import Path

import numpy as np
import pandas as pd
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.lines import Line2D
from matplotlib.patches import Patch
plt.rcParams["font.family"] = "DejaVu Sans"
plt.rcParams["axes.unicode_minus"] = False

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ======================================================================
# ТОХИРГОО
# ======================================================================
BASE_DIR = Path(__file__).parent
ARCHIVE_DIR = BASE_DIR / "archive"
OUTPUT_DIR = BASE_DIR / "reports"
UPLOAD_DIR = BASE_DIR / "dashboard_uploads"

STREAMLIT_URL = "https://loan-dashboardgit-jht3b94lsjnllhvcdu7swo.streamlit.app/"

RED, GREEN, BLUE = "#e24b4a", "#1d9e75", "#2563eb"
PURPLE, ORANGE = "#7c3aed", "#f59e0b"
BUCKET_COLS = ["#84cc16", "#f59e0b", "#f97316", "#ef4444", "#dc2626", "#7f1d1d"]
# ======================================================================


def load_final_data():
    files = glob.glob(str(ARCHIVE_DIR / "*.parquet"))
    if not files:
        raise FileNotFoundError(f"{ARCHIVE_DIR} дотроос parquet олдсонгүй. Эхлээд etl_cowork.py-г ажиллуулна уу.")
    path = max(files, key=os.path.getmtime)
    df = pd.read_parquet(path)
    print(f"Эцсийн дата уншлаа: {os.path.basename(path)} ({len(df):,} мөр)")
    return df, os.path.basename(path)


def export_for_dashboard(df: pd.DataFrame) -> Path:
    UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
    today = dt.date.today().strftime("%Y-%m-%d")
    out = UPLOAD_DIR / f"dashboard_upload_{today}.csv"
    df.to_csv(out, index=False, encoding="utf-8-sig")
    print(f"Дашбордын upload файл хадгалагдлаа: {out.name}")
    return out


# ── Тооцоолол ───────────────────────────────────────────────────────────────
def _b2n(s):
    return s.map(lambda v: 1.0 if v in (True, 1, 1.0) else (0.0 if v in (False, 0, 0.0) else np.nan))


def _mean(s):
    return pd.to_numeric(s, errors="coerce").mean()


def compute(df: pd.DataFrame) -> dict:
    S = {}
    OD = "max_active_overdue_day"
    S["n_loan"] = len(df)
    vc = df["status_1"].value_counts()
    S["nC"] = int(vc.get("C", 0)); S["nOmax"] = int(vc.get("O_max", 0)); S["nOact"] = int(vc.get("O_active", 0))
    S["nOpen"] = S["nOmax"] + S["nOact"]
    oa = df[df["status_1"] == "O_active"]; oad = oa[OD]
    S["oa_mean"] = oad.mean(); S["oa_med"] = oad.median(); S["oa_max"] = oad.max()
    S["buckets"] = [("1", int(oad.between(1, 1).sum())), ("2-5", int(oad.between(2, 5).sum())),
                    ("6-10", int(oad.between(6, 10).sum())), ("11-15", int(oad.between(11, 15).sum())),
                    ("16-30", int(oad.between(16, 30).sum())), ("30+", int(oad.between(31, 9999).sum()))]
    g = df.groupby("cust_code")
    cust = g.agg(max_aod=(OD, "max")).reset_index()
    attrs = ["age", "gender", "fin_score", "psy_score", "total_score", "total_score_sr",
             "slry_last_amt", "slry_last_avg_6m", "slry_last_row_cnt_24m", "zms_active_ln_cnt",
             "zms_monthly_payment", "has_ios", "is_bio_login", "is_device_remember", "mobile_no",
             "slry_has_cont_salary_3m"]
    attrs = [c for c in attrs if c in df.columns]
    cust = cust.merge(g[attrs].first().reset_index(), on="cust_code")
    cust["has_od"] = cust["max_aod"] > 0
    S["n_cust"] = len(cust)
    S["od_cust"] = int(cust["has_od"].sum()); S["od_cust_pct"] = cust["has_od"].mean() * 100
    S["c30"] = int((cust["max_aod"] > 30).sum()); S["c30_pct"] = (cust["max_aod"] > 30).mean() * 100
    S["age_mean"] = _mean(cust["age"])
    for c in ["total_score", "fin_score", "psy_score"]:
        s = pd.to_numeric(cust[c], errors="coerce")
        S[c + "_mean"] = s.mean(); S[c + "_min"] = s.min(); S[c + "_max"] = s.max(); S[c + "_med"] = s.median()
    S["corr"] = oa[["total_score", OD]].dropna().corr().iloc[0, 1] if len(oa) > 2 else float("nan")
    oa15 = oa[oa[OD] >= 15]; oa30 = oa[oa[OD] >= 30]
    S["corr15"] = oa15[["total_score", OD]].dropna().corr().iloc[0, 1] if len(oa15) > 2 else float("nan")
    S["corr30"] = oa30[["total_score", OD]].dropna().corr().iloc[0, 1] if len(oa30) > 2 else float("nan")
    S["corr15_n"] = len(oa15); S["corr30_n"] = len(oa30)
    od = cust[cust["has_od"]]; nod = cust[~cust["has_od"]]
    S["cmp"] = {}
    for c in ["total_score", "fin_score", "psy_score", "age", "slry_last_amt", "slry_last_avg_6m", "zms_active_ln_cnt", "zms_monthly_payment"]:
        if c in cust.columns:
            S["cmp"][c] = (_mean(od[c]), _mean(nod[c]))
    S["bcmp"] = {}
    for c in ["has_ios", "is_bio_login", "slry_has_cont_salary_3m", "mobile_no", "gender"]:
        if c in cust.columns:
            S["bcmp"][c] = (_b2n(od[c]).mean() * 100, _b2n(nod[c]).mean() * 100)
    S["ios_all"] = _b2n(cust["has_ios"]).mean() * 100 if "has_ios" in cust else float("nan")
    S["bio_all"] = _b2n(cust["is_bio_login"]).mean() * 100 if "is_bio_login" in cust else float("nan")
    S["male_all"] = _b2n(cust["gender"]).mean() * 100 if "gender" in cust else float("nan")
    cust["ag"] = pd.cut(pd.to_numeric(cust["age"], errors="coerce"), bins=[0, 29, 45, 200], labels=["17-29", "30-45", "46+"])
    cust["g"] = _b2n(cust["gender"]).map({1.0: "Эрэгтэй", 0.0: "Эмэгтэй"})
    segs = []
    for a in ["17-29", "30-45", "46+"]:
        for gg in ["Эрэгтэй", "Эмэгтэй"]:
            sub = cust[(cust["ag"].astype(str) == a) & (cust["g"] == gg)]
            if len(sub) == 0:
                continue
            segs.append({"name": f"{a} {gg}", "age": a, "g": gg, "n": len(sub),
                         "pct": len(sub) / len(cust) * 100, "od": sub["has_od"].mean() * 100,
                         "p30": (sub["max_aod"] > 30).mean() * 100, "tot": _mean(sub["total_score"]),
                         "fin": _mean(sub["fin_score"]), "psy": _mean(sub["psy_score"]),
                         "sal": _mean(sub["slry_last_amt"]) / 1e6})
    S["segs"] = segs
    S["seg_hi"] = max(segs, key=lambda r: r["od"]); S["seg_lo"] = min(segs, key=lambda r: r["od"])
    S["by_g"] = {gg: (len(cust[cust["g"] == gg]), cust[cust["g"] == gg]["has_od"].mean() * 100) for gg in ["Эрэгтэй", "Эмэгтэй"]}
    S["by_age"] = {a: (len(cust[cust["ag"].astype(str) == a]), cust[cust["ag"].astype(str) == a]["has_od"].mean() * 100) for a in ["17-29", "30-45", "46+"]}
    S["_cust"] = cust
    return S


# ── Динамик тайлбарын туслахууд ───────────────────────────────────────────────
def _assess_od(p):
    if p < 8:  return "харьцангуй бага, багцын чанар сайн түвшинд"
    if p < 13: return "дунд зэргийн түвшинд"
    if p < 20: return "нэлээд өндөр, анхаарал шаардсан түвшинд"
    return "өндөр, эрсдэл ихтэй түвшинд"


def _assess_c30(p):
    if p < 1:  return "маш бага"
    if p < 3:  return "бага"
    if p < 6:  return "анхаарал татахуйц"
    return "өндөр"


def _assess_corr(r):
    a = abs(r)
    if a < 0.05: return "маш сул"
    if a < 0.15: return "сул"
    if a < 0.30: return "дунд зэргийн"
    return "харьцангуй хүчтэй"


def _assess_short(p):
    if p >= 60: return "дийлэнх нь богино (≤5 хоног) хэтрэлт бөгөөд эргэн төлөлтийн зан төлөв эерэг"
    if p >= 45: return "ихэнх нь богино хугацааны хэтрэлт"
    return "богино болон урт хугацааны хэтрэлт харьцангуй жигд тархсан"


def _trend(cur, prev, unit="пп"):
    d = cur - prev
    if abs(d) < 0.3:
        return f"өмнөх үетэй харьцуулахад бараг тогтвортой ({prev:.1f} → {cur:.1f})"
    return f"өмнөх үеэс {abs(d):.1f}{unit}-ээр {'нэмэгдсэн' if d > 0 else 'буурсан'} ({prev:.1f} → {cur:.1f})"


def _best_discriminator(tspr, fspr, pspr):
    return max({"нийт оноо": tspr, "санхүүгийн оноо": fspr, "сэтгэлзүйн оноо": pspr}.items(), key=lambda kv: kv[1])[0]


def _prev_summary(current_filename):
    """archive дахь өмнөх үеийн (хоёр дахь шинэ) parquet-аас гол үзүүлэлт авна. Үгүй бол None."""
    files = [f for f in glob.glob(str(ARCHIVE_DIR / "*.parquet")) if os.path.basename(f) != current_filename]
    if not files:
        return None
    path = max(files, key=os.path.getmtime)
    try:
        dfp = pd.read_parquet(path)
        g = dfp.groupby("cust_code"); cu = g.agg(m=("max_active_overdue_day", "max")).reset_index()
        return {"name": os.path.basename(path).replace(".parquet", ""),
                "od_pct": (cu["m"] > 0).mean() * 100,
                "c30_pct": (cu["m"] > 30).mean() * 100,
                "score": pd.to_numeric(g["total_score"].first(), errors="coerce").mean()}
    except Exception:
        return None


# ── Графикууд ─────────────────────────────────────────────────────────────────
def _save(fig):
    b = io.BytesIO(); fig.tight_layout(); fig.savefig(b, format="png", dpi=150); plt.close(fig); b.seek(0)
    return b


def _band_chart(cust, col, title, xlabel, step=30):
    s = pd.to_numeric(cust[col], errors="coerce")
    lo = int(np.floor(s.min() / step) * step); hi = int(np.ceil(s.max() / step) * step)
    edges = list(range(lo, hi + step, step))
    labels = [f"{edges[i]}-{edges[i + 1]}" for i in range(len(edges) - 1)]
    d = cust.copy(); d["_bb"] = pd.cut(s, bins=edges, labels=labels, include_lowest=True, duplicates="drop")
    d["_od"] = d["max_aod"] > 0; d["_od15"] = d["max_aod"] >= 15
    t = d.groupby("_bb", observed=True).agg(n=("_od", "count"), od=("_od", "mean"), od15=("_od15", "mean")).reset_index()
    t["od"] *= 100; t["od15"] *= 100
    fig, ax = plt.subplots(figsize=(8, 3.7))
    ax.bar(t["_bb"].astype(str), t["n"], color=BLUE, alpha=.45, label="Харилцагч (тоо)")
    ax2 = ax.twinx()
    ax2.plot(t["_bb"].astype(str), t["od"], color=RED, marker="o", lw=2, label="Хэтрэлт % (1+ хоног)")
    ax2.plot(t["_bb"].astype(str), t["od15"], color=PURPLE, marker="s", lw=2, ls="--", label="Хэтрэлт % (15+ хоног)")
    for xi, yi in zip(range(len(t)), t["od"]):
        ax2.annotate(f"{yi:.1f}%", (xi, yi), ha="center", va="bottom", color=RED, fontsize=7)
    for xi, yi in zip(range(len(t)), t["od15"]):
        ax2.annotate(f"{yi:.1f}%", (xi, yi), ha="center", va="top", color=PURPLE, fontsize=7)
    ax.set_title(title); ax.set_xlabel(xlabel); ax.set_ylabel("Харилцагч"); ax2.set_ylabel("Хэтрэлт %")
    plt.setp(ax.get_xticklabels(), rotation=12, ha="right")
    h1, l1 = ax.get_legend_handles_labels(); h2, l2 = ax2.get_legend_handles_labels()
    ax2.legend(h1 + h2, l1 + l2, fontsize=7, loc="upper right")
    return _save(fig)


def make_charts(S):
    cust = S["_cust"]; C = {}
    fig, ax = plt.subplots(figsize=(8, 3.4))
    ax.hist(pd.to_numeric(cust["total_score"], errors="coerce").dropna(), bins=40, color=BLUE, alpha=.85)
    ax.axvline(S["total_score_mean"], color=RED, ls="--", lw=1.5, label=f"Дундаж {S['total_score_mean']:.0f}")
    ax.set_title("Нийт онооны тархалт (харилцагч)"); ax.set_xlabel("Нийт оноо"); ax.set_ylabel("Тоо"); ax.legend()
    C["dist"] = _save(fig)
    fig, ax = plt.subplots(figsize=(8, 3.4))
    labs = [b[0] for b in S["buckets"]]; vals = [b[1] for b in S["buckets"]]
    bars = ax.bar(labs, vals, color=BUCKET_COLS)
    for r, v in zip(bars, vals):
        ax.annotate(f"{v:,}", (r.get_x() + r.get_width() / 2, v), ha="center", va="bottom", fontsize=9)
    ax.set_title("Идэвхтэй хэтрэлтийн бүс (данс, хоногоор)"); ax.set_ylabel("Дансны тоо"); ax.set_xlabel("Хэтрэлтийн хоног")
    C["bucket"] = _save(fig)
    C["band_total"] = _band_chart(cust, "total_score", "Нийт онооны бүсээр — харилцагчийн тоо ба хэтрэлтийн хувь", "Нийт онооны бүс")
    C["band_fin"] = _band_chart(cust, "fin_score", "Санхүүгийн онооны бүсээр — харилцагчийн тоо ба хэтрэлтийн хувь", "Санхүүгийн онооны бүс")
    C["band_psy"] = _band_chart(cust, "psy_score", "Сэтгэлзүйн онооны бүсээр — харилцагчийн тоо ба хэтрэлтийн хувь", "Сэтгэлзүйн онооны бүс")
    fig, ax = plt.subplots(figsize=(8, 3.7))
    mm = [("total_score", "Нийт оноо"), ("fin_score", "Санхүүгийн"), ("psy_score", "Сэтгэлзүйн")]
    g0 = cust[cust["max_aod"] == 0]; g1 = cust[cust["max_aod"] > 0]; g15 = cust[cust["max_aod"] >= 15]
    groups = [("Хэтрээгүй", g0, GREEN), ("Хэтэрсэн (1+)", g1, ORANGE), ("15+ хэтэрсэн", g15, RED)]
    x = np.arange(len(mm)); w = .26
    for i, (nm, gd, col) in enumerate(groups):
        vals = [pd.to_numeric(gd[c], errors="coerce").mean() for c, _ in mm]
        bb = ax.bar(x + (i - 1) * w, vals, w, label=nm, color=col)
        for r in bb:
            ax.annotate(f"{r.get_height():.0f}", (r.get_x() + r.get_width() / 2, r.get_height()), ha="center", va="bottom", fontsize=7)
    ax.set_xticks(x); ax.set_xticklabels([t for _, t in mm]); ax.set_title("Дундаж оноо: хэтрээгүй / хэтэрсэн (1+) / 15+ хэтэрсэн"); ax.set_ylabel("Оноо"); ax.legend(fontsize=8)
    C["od_score"] = _save(fig)
    cust["seg"] = cust["ag"].astype(str) + " " + cust["g"].astype(str)
    order = [s["name"] for s in S["segs"]]
    none_ = [(cust[cust["seg"] == nm]["max_aod"] == 0).mean() * 100 for nm in order]
    od_ = [(cust[cust["seg"] == nm]["max_aod"] > 0).mean() * 100 for nm in order]
    od15_ = [(cust[cust["seg"] == nm]["max_aod"] >= 15).mean() * 100 for nm in order]
    x = np.arange(len(order)); w = .26
    fig, ax = plt.subplots(figsize=(9.2, 3.9))
    for off, vals, nm, col in [(-w, none_, "Хэтрээгүй %", GREEN), (0, od_, "Хэтэрсэн % (1+)", ORANGE), (w, od15_, "15+ хэтэрсэн %", RED)]:
        bb = ax.bar(x + off, vals, w, label=nm, color=col)
        for r in bb:
            ax.annotate(f"{r.get_height():.1f}", (r.get_x() + r.get_width() / 2, r.get_height()), ha="center", va="bottom", fontsize=7)
    ax.set_xticks(x); ax.set_xticklabels(order, rotation=15, ha="right"); ax.set_title("Сегмент бүрийн хэтрэлтийн бүтэц (хэтрээгүй / 1+ / 15+)"); ax.set_ylabel("Хувь %"); ax.legend(fontsize=8)
    C["seg_od"] = _save(fig)
    fig, ax = plt.subplots(figsize=(8, 3.6))
    names = [s["name"] for s in S["segs"]]; sizes = [s["n"] for s in S["segs"]]
    bars = ax.barh(names[::-1], sizes[::-1], color=BUCKET_COLS[:len(names)][::-1])
    for r, v in zip(bars, sizes[::-1]):
        ax.annotate(f"{v:,}", (v, r.get_y() + r.get_height() / 2), va="center", fontsize=8)
    ax.set_title("Сегментийн хэмжээ (харилцагчийн тоо)"); ax.set_xlabel("Харилцагч")
    C["seg_size"] = _save(fig)
    return C


# ── DOCX туслахууд ────────────────────────────────────────────────────────────
def _shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), fill); tcPr.append(shd)


def _setw(cell, dxa):
    tcPr = cell._tc.get_or_add_tcPr(); tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(dxa)); tcW.set(qn("w:type"), "dxa"); tcPr.append(tcW)


def _pagenum(par):
    r = par.add_run(); f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = "PAGE"
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "end")
    r._r.append(f1); r._r.append(it); r._r.append(f2)


def _H1(doc, t):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(8)
    r = p.add_run(t); r.bold = True; r.font.size = Pt(14); r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)


def _BODY(doc, t):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.space_after = Pt(8)
    p.add_run(t).font.size = Pt(11)


def _CAP(doc, t):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(10)
    r = p.add_run(t); r.italic = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def _BUL(doc, t):
    p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(3)
    p.add_run(t).font.size = Pt(11)


def _IMG(doc, buf, w=6.2):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(buf, width=Inches(w))


def _TABLE(doc, widths, header, rows):
    t = doc.add_table(rows=1, cols=len(header)); t.style = "Table Grid"; t.alignment = 1
    for i, h in enumerate(header):
        c = t.rows[0].cells[i]; _setw(c, widths[i]); _shade(c, "D5E8F0")
        c.paragraphs[0].alignment = (WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER)
        run = c.paragraphs[0].add_run(h); run.bold = True; run.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            _setw(cells[i], widths[i])
            cells[i].paragraphs[0].alignment = (WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER)
            cells[i].paragraphs[0].add_run(str(v)).font.size = Pt(9.5)


def _m(x):
    try:
        return f"{float(x):,.0f}"
    except (TypeError, ValueError):
        return "—"


# ── Хавсралт: хамаарлын хүснэгт ───────────────────────────────────────────────
_VARS = {"fin_score": "Санх.оноо", "psy_score": "Сэтгэлз.оноо", "total_score": "Нийт оноо",
         "slry_last_amt": "Цалин", "slry_last_avg_6m": "Цалин6с", "slry_last_row_cnt_24m": "Цалин.бичилт",
         "zms_active_ln_cnt": "ЗМС зээл", "zms_monthly_payment": "ЗМС төлбөр", "slry_has_cont_salary_3m": "Тас.цалин",
         "has_ios": "iOS", "is_bio_login": "Биометр", "is_device_remember": "Төх.санасан", "mobile_no": "88/99утас", "age": "Нас"}
_BNAMES = ["0", "1-15", "16-30", "30+", "Нийт"]


def _bf(a, bn):
    if bn == "0": return (a == 0).astype(float)
    if bn == "1-15": return a.between(1, 15).astype(float)
    if bn == "16-30": return a.between(16, 30).astype(float)
    if bn == "30+": return (a > 30).astype(float)
    return a.astype(float)


def _add_corr_appendix(doc, S):
    cc = S["_cust"].copy(); cc["seg"] = cc["ag"].astype(str) + " " + cc["g"].astype(str)
    X = pd.DataFrame({k: (_b2n(cc[k]) if cc[k].dtype == object or cc[k].dtype == bool else pd.to_numeric(cc[k], errors="coerce"))
                      for k in _VARS if k in cc.columns})

    def top5(idx, tgt):
        Xs = X.loc[idx]; r = {}
        for k in X.columns:
            v = Xs[k]
            if v.notna().sum() < 5 or v.std() == 0 or tgt.std() == 0:
                continue
            val = np.corrcoef(v.fillna(v.mean()), tgt)[0, 1]
            if not np.isnan(val):
                r[k] = val
        return [f"{_VARS[k]}{'↑' if val > 0 else '↓'}{abs(val):.2f}" for k, val in sorted(r.items(), key=lambda kv: abs(kv[1]), reverse=True)[:5]]

    order = [s["name"] for s in S["segs"]]; segn = {s["name"]: s["n"] for s in S["segs"]}
    sec = doc.add_section(WD_SECTION.NEW_PAGE); sec.orientation = WD_ORIENT.LANDSCAPE
    _w, _h = sec.page_width, sec.page_height; sec.page_width = _h; sec.page_height = _w
    sec.left_margin = Inches(0.5); sec.right_margin = Inches(0.5); sec.top_margin = Inches(0.6); sec.bottom_margin = Inches(0.6)
    _H1(doc, "Хавсралт. Сегмент × хэтрэлтийн бүс — хамаарлын шинжилгээ (top 5)")
    _BODY(doc, "Нүд бүрд тухайн сегментийн харилцагчид тухайн хэтрэлтийн бүст багтахтай хамгийн өндөр хамааралтай (|корреляц|) 5 хувьсагчийг чиглэл (↑ өндөр утга, ↓ бага утга) болон коэффициентийн хамт жагсаав. 'Нийт' багана нь хэтрэлтийн хоногтой шууд хамаарлыг харуулна. Цөөн ажиглалттай (n бага) нүднүүд статистик найдвар багатайг анхаарна уу.")
    widths = [1900, 2442, 2442, 2442, 2442, 2442]; header = ["Сегмент (n)"] + _BNAMES
    tb = doc.add_table(rows=1, cols=6); tb.style = "Table Grid"
    for i, ht in enumerate(header):
        cl = tb.rows[0].cells[i]; _setw(cl, widths[i]); _shade(cl, "D5E8F0"); cl.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = cl.paragraphs[0].add_run(ht); rr.bold = True; rr.font.size = Pt(8.5)
    for seg in order:
        sub = cc[cc["seg"] == seg]; a = sub["max_aod"]
        cells = tb.add_row().cells
        c0 = cells[0]; _setw(c0, widths[0]); r0 = c0.paragraphs[0].add_run(seg); r0.bold = True; r0.font.size = Pt(8)
        c0.paragraphs[0].add_run().add_break(); rn = c0.paragraphs[0].add_run(f"n={segn[seg]:,}"); rn.font.size = Pt(7); rn.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        for j, bn in enumerate(_BNAMES):
            cl = cells[j + 1]; _setw(cl, widths[j + 1]); p = cl.paragraphs[0]
            tgt = _bf(a, bn); n_in = int((tgt > 0).sum()) if bn != "Нийт" else len(sub)
            rn = p.add_run(f"n={n_in:,}"); rn.font.size = Pt(6.5); rn.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            lines = top5(sub.index, tgt)
            if not lines:
                p.add_run().add_break(); p.add_run("—").font.size = Pt(7.5)
            else:
                for ln in lines:
                    p.add_run().add_break(); p.add_run(ln).font.size = Pt(7.5)


def _add_group_profile(doc, S):
    """Хэтрээгүй / +1 / +30 бүлгийн ерөнхий зан төлөвийн харьцуулалт (динамик тайлбар)."""
    cu = S["_cust"]
    G0 = cu[cu["max_aod"] == 0]; G1 = cu[cu["max_aod"] >= 1]; G30 = cu[cu["max_aod"] >= 30]

    def mn(G, c):
        return pd.to_numeric(G[c], errors="coerce").mean() if c in G.columns else float("nan")

    def pc(G, c):
        return _b2n(G[c]).mean() * 100 if c in G.columns else float("nan")

    od = (cu["max_aod"] > 0).astype(float)
    z = pd.to_numeric(cu["zms_active_ln_cnt"], errors="coerce") if "zms_active_ln_cnt" in cu.columns else None
    zcorr = np.corrcoef(z.fillna(z.mean()), od)[0, 1] if (z is not None and z.std() > 0) else float("nan")

    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    sec.orientation = WD_ORIENT.PORTRAIT
    sec.page_width = Inches(8.5); sec.page_height = Inches(11)
    sec.left_margin = Inches(1); sec.right_margin = Inches(1); sec.top_margin = Inches(1); sec.bottom_margin = Inches(1)

    _H1(doc, "Хавсралт-2. Хэтрээгүй / +1 / +30 бүлгийн ерөнхий зан төлөвийн харьцуулалт")
    _BODY(doc, f"Хэтрэлт хүндрэх тусам санхүүгийн чадавхын бараг бүх үзүүлэлт тогтмол буурах хандлагатай байна: нийт оноо {mn(G0,'total_score'):.0f} → {mn(G1,'total_score'):.0f} → {mn(G30,'total_score'):.0f}, санхүүгийн оноо {mn(G0,'fin_score'):.0f} → {mn(G1,'fin_score'):.0f} → {mn(G30,'fin_score'):.0f}, сүүлийн цалин {mn(G0,'slry_last_amt')/1e6:.2f} → {mn(G1,'slry_last_amt')/1e6:.2f} → {mn(G30,'slry_last_amt')/1e6:.2f} сая ₮.")
    _BODY(doc, f"Хамгийн тод ялгарах хоёр шинж нь: (1) сүүлийн 3 сарын тасралтгүй цалинтай харилцагчийн хувь хэтрээгүй бүлэгт {pc(G0,'slry_has_cont_salary_3m'):.1f}% байсан бол +30 бүлэгт {pc(G30,'slry_has_cont_salary_3m'):.1f}% болж огцом буурсан нь орлогын тогтворгүй байдал өндөр эрсдэлтэй шууд холбоотой; (2) эрэгтэй харилцагчийн эзлэх хувь {pc(G0,'gender'):.1f}% → {pc(G1,'gender'):.1f}% → {pc(G30,'gender'):.1f}% болж огцом нэмэгдсэн нь эрэгтэйчүүд эрсдэл өндөртэйг харуулна. Биометр нэвтрэлт идэвхжүүлсэн хувь хэтэрсэн бүлгүүдэд бага ({pc(G0,'is_bio_login'):.1f}% → {pc(G1,'is_bio_login'):.1f}% → {pc(G30,'is_bio_login'):.1f}%).")
    _BODY(doc, f"Бусад зээлийн (ЗМС) мэдээллийн хувьд хэтрэлт хүндрэх тусам идэвхтэй зээлийн тоо ({mn(G0,'zms_active_ln_cnt'):.1f} → {mn(G30,'zms_active_ln_cnt'):.1f}), сарын төлбөр ({mn(G0,'zms_monthly_payment')/1e3:.0f} → {mn(G30,'zms_monthly_payment')/1e3:.0f} мянга) болон хаагдсан зээлийн нийт дүн ({mn(G0,'zms_closed_ln_total_amount')/1e6:.1f} → {mn(G30,'zms_closed_ln_total_amount')/1e6:.1f} сая ₮) бүгд буурч, хэтрэлттэй корреляц {zcorr:+.2f} орчим сул сөрөг байна. Энэ нь хэтэрсэн харилцагчид бусад зээлээр хэт ачаалалтай байгаагаас бус, харин зээлийн түүх нимгэн, тогтсон зээлийн харилцаа сул байгааг илтгэнэ — баялаг зээлийн түүхтэй харилцагчид эрсдэл бага байна.")
    _BODY(doc, f"Сонирхолтой нь сэтгэлзүйн оноо хэтрэлт хүндрэх тусам буурахгүй, бүр бага зэрэг нэмэгдэж ({mn(G0,'psy_score'):.1f} → {mn(G30,'psy_score'):.1f}) байгаа нь сэтгэлзүйн оноо хэтрэлтийг сул ялгадаг өмнөх дүгнэлтийг баталж байна.")


def _seg_grid(cu, order):
    scores = [("total_score", "Нийт оноо", 330, 570), ("fin_score", "Санхүүгийн оноо", 150, 330), ("psy_score", "Сэтгэлзүйн оноо", 90, 270)]
    fig, axes = plt.subplots(len(order), 3, figsize=(13, 16))
    for i, seg in enumerate(order):
        sub = cu[cu["seg"] == seg]
        for j, (col, nm, lo, hi) in enumerate(scores):
            ax = axes[i, j]
            edges = list(range(lo, hi + 30, 30)); labels = [f"{edges[k]}-{edges[k+1]}" for k in range(len(edges) - 1)]
            s = pd.to_numeric(sub[col], errors="coerce")
            b = pd.cut(s, bins=edges, labels=labels, include_lowest=True)
            t = sub.assign(_b=b).groupby("_b", observed=False).agg(
                n=("max_aod", "count"),
                od1=("max_aod", lambda x: (x > 0).mean() * 100 if len(x) else 0),
                od15=("max_aod", lambda x: (x >= 15).mean() * 100 if len(x) else 0)).reindex(labels)
            x = np.arange(len(labels))
            ax.bar(x, t["n"].fillna(0).values, color=BLUE, alpha=.35)
            ax2 = ax.twinx()
            ax2.plot(x, t["od1"].values, color=RED, marker="o", ms=3, lw=1.4)
            ax2.plot(x, t["od15"].values, color=PURPLE, marker="s", ms=3, lw=1.4, ls="--")
            ax2.set_ylim(0, 30)
            ax.set_xticks(x); ax.set_xticklabels([l.split("-")[0] for l in labels], fontsize=6); ax.tick_params(labelsize=6)
            ax2.tick_params(labelsize=6, colors=RED)
            if i == 0: ax.set_title(nm, fontsize=11, fontweight="bold")
            if j == 0: ax.set_ylabel(seg, fontsize=8, fontweight="bold")
    leg = [Patch(facecolor=BLUE, alpha=.35, label="Харилцагч (тоо)"),
           Line2D([0], [0], color=RED, marker="o", label="Хэтрэлт % (1+)"),
           Line2D([0], [0], color=PURPLE, marker="s", ls="--", label="Хэтрэлт % (15+)")]
    fig.legend(handles=leg, loc="upper center", ncol=3, fontsize=10, bbox_to_anchor=(0.5, 0.985))
    fig.suptitle("Сегмент × онооны төрлөөр — онооны histogram, 1+ ба 15+ хэтрэлт %", fontsize=13, y=0.998)
    fig.tight_layout(rect=[0, 0, 1, 0.965])
    return _save(fig)


def _add_six_segments(doc, S):
    cu = S["_cust"].copy(); cu["seg"] = cu["ag"].astype(str) + " " + cu["g"].astype(str)
    order = [s["name"] for s in S["segs"]]
    buf = _seg_grid(cu, order)
    sec = doc.add_section(WD_SECTION.NEW_PAGE); sec.orientation = WD_ORIENT.PORTRAIT
    sec.page_width = Inches(8.5); sec.page_height = Inches(11)
    sec.left_margin = Inches(0.8); sec.right_margin = Inches(0.8); sec.top_margin = Inches(0.8); sec.bottom_margin = Inches(0.8)
    _H1(doc, "Хавсралт-3. 6 сегментийн хэтрэлтийн бүтэц ба онооны хамаарал")
    rows = []
    for seg in order:
        a = cu[cu["seg"] == seg]["max_aod"]
        rows.append([seg, _m(len(a)), f"{(a == 0).mean()*100:.1f}", f"{(a > 0).mean()*100:.1f}", f"{(a >= 15).mean()*100:.1f}", f"{(a > 30).mean()*100:.1f}"])
    _TABLE(doc, [2640, 1140, 1440, 1080, 1080, 980], ["Сегмент", "n", "Хэтрээгүй %", "1+ %", "15+ %", "30+ %"], rows)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    hi = S["seg_hi"]; lo = S["seg_lo"]
    _BODY(doc, f"Дээрх хүснэгтэд харилцагчдыг нас, хүйсээр зургаан сегментэд хувааж, хэтрэлтийн бүтцийг (хэтрээгүй / 1+ / 15+ / 30+ хоног) харуулав. {hi['name']} сегмент бүх түвшинд хамгийн өндөр эрсдэлтэй (1+ хэтрэлт {hi['od']:.1f}%, 30+ {hi['p30']:.2f}%) бол {lo['name']} хамгийн бага ({lo['od']:.1f}%) байна.")
    _IMG(doc, buf, w=6.8); _CAP(doc, "Зураг 9. Сегмент (мөр) × онооны төрөл (багана) — хэтрэлт % (1+) онооны бүсээр")
    _BODY(doc, "Графикийн торноос харахад ихэнх сегментэд нийт болон санхүүгийн оноо нэмэгдэх тусам хэтрэлт буурах тод хандлага ажиглагдана. Сэтгэлзүйн онооны хувьд энэ хандлага сул буюу зарим сегментэд бараг хавтгай байгаа нь сэтгэлзүйн оноо хэтрэлтийг сул ялгадаг өмнөх дүгнэлттэй нийцэж байна.")


# ── Тайлан угсрах ─────────────────────────────────────────────────────────────
def build_report(S, src_name, outdir: Path) -> Path:
    outdir.mkdir(parents=True, exist_ok=True)
    today = dt.date.today().strftime("%Y.%m.%d")
    doc = Document()
    doc.styles["Normal"].font.name = "Arial"; doc.styles["Normal"].font.size = Pt(11)

    fp = doc.sections[0].footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run(f"PsyFint дүн шинжилгээ · {today} · "); fr.font.size = Pt(8); fr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    _pagenum(fp)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PsyFint кредит скоринг — харилцагчийн сегмент ба хугацаа хэтрэлтийн дүн шинжилгээ"); r.bold = True; r.font.size = Pt(15)
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER; p2.paragraph_format.space_after = Pt(12)
    r2 = p2.add_run(f"Тайлан гаргасан огноо: {today}  ·  Эх өгөгдөл: {src_name}"); r2.font.size = Pt(10); r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    C = make_charts(S)
    # нэмэлт тоонууд
    cu = S["_cust"]
    fem = 100 - S["male_all"]
    b17n, b17o = S["by_age"]["17-29"]; b30n, b30o = S["by_age"]["30-45"]; b46n, b46o = S["by_age"]["46+"]
    em = S["by_g"]["Эрэгтэй"]; ef = S["by_g"]["Эмэгтэй"]
    hi = S["seg_hi"]; lo = S["seg_lo"]
    bio = S["bcmp"].get("is_bio_login"); cont = S["bcmp"].get("slry_has_cont_salary_3m"); male = S["bcmp"].get("gender")
    seg46m = next((s for s in S["segs"] if s["name"] == "46+ Эрэгтэй"), {"n": 0})
    _bk = dict(S["buckets"]); _tot = max(sum(v for _, v in S["buckets"]), 1)
    SHORT = _bk["1"] + _bk["2-5"]; MID = _bk["6-10"] + _bk["11-15"]; LONG = _bk["16-30"] + _bk["30+"]
    SHORT_P = SHORT / _tot * 100; MID_P = MID / _tot * 100; LONG_P = LONG / _tot * 100

    def _br(col, step=30):
        s = pd.to_numeric(cu[col], errors="coerce"); lo_ = int(np.floor(s.min() / step) * step); hi_ = int(np.ceil(s.max() / step) * step)
        edges = list(range(lo_, hi_ + step, step)); labels = [f"{edges[i]}-{edges[i + 1]}" for i in range(len(edges) - 1)]
        d = cu.copy(); d["b"] = pd.cut(s, bins=edges, labels=labels, include_lowest=True)
        t = d.groupby("b", observed=True).agg(n=("has_od", "count"), od=("has_od", "mean")).reset_index(); t["od"] *= 100
        t = t[t["n"] >= 20]
        return (t.iloc[0]["b"], t.iloc[0]["od"], t.iloc[-1]["b"], t.iloc[-1]["od"], t["od"].max() - t["od"].min())
    tlo_b, tlo, thi_b, thi, tspr = _br("total_score"); fspr = _br("fin_score")[4]; pspr = _br("psy_score")[4]
    prev = _prev_summary(src_name)

    _H1(doc, "Тойм")
    _BODY(doc, f"Энэхүү тайлан нь PsyFint кредит скорингийн загвараар үнэлэгдэж зээл олгогдсон харилцагчдын зээлийн гүйцэтгэл болон хугацаа хэтрэлтийн эрсдэлийг иж бүрэн дүгнэх зорилготой. Шинжилгээг хоёр түвшинд гүйцэтгэв: зээл (данс) түвшинд тус бүр олгосон зээлийн төлөв, хэтрэлтийг, харин харилцагч түвшинд нэг хүний бүх зээлийг нэгтгэн авч үзэв. Нийт {_m(S['n_cust'])} харилцагчид олгогдсон {_m(S['n_loan'])} зээл хамрагдсан бөгөөд нэг харилцагч дунджаар хэд хэдэн зээлтэй байгаа нь давтан зээлийн идэвхжлийг харуулна.")
    _BODY(doc, f"Багцын ерөнхий чанарын хувьд нийт зээлийн {S['nC']/S['n_loan']*100:.1f}% ({_m(S['nC'])}) аль хэдийн хаагдсан, {S['nOpen']/S['n_loan']*100:.1f}% ({_m(S['nOpen'])}) нээлттэй хэвээр байгаагаас {S['nOact']/S['n_loan']*100:.1f}% ({_m(S['nOact'])}) нь идэвхтэй хугацаа хэтрэлтэд орсон байна. Харилцагч түвшинд дүр зураг арай өндөр буюу нийт харилцагчдын {S['od_cust_pct']:.1f}% ({_m(S['od_cust'])}) ямар нэг идэвхтэй хэтрэлттэй байгаа нь нэг харилцагчийн хэд хэдэн зээлийн аль нэгэнд хэтрэлт гарахад тухайн харилцагч бүхэлдээ хэтрэлттэй гэж тооцогддогтой холбоотой. Үүнээс {S['c30_pct']:.2f}% ({_m(S['c30'])}) нь 30-аас дээш хоногийн хэтрэлттэй буюу шууд анхаарал шаардсан өндөр эрсдэлийн бүсэд байна.")
    _t = f"Идэвхтэй хэтрэлтийн ерөнхий түвшин {_assess_od(S['od_cust_pct'])} байгаа бөгөөд 30+ хоногийн эрсдэл {_assess_c30(S['c30_pct'])} байна."
    if prev:
        _t += f" Өмнөх үетэй ({prev['name']}) харьцуулахад идэвхтэй хэтрэлт {_trend(S['od_cust_pct'], prev['od_pct'])}, 30+ хоногийн эрсдэл {_trend(S['c30_pct'], prev['c30_pct'])}."
    _BODY(doc, _t)
    _BODY(doc, "Дараах хэсгүүдэд онооны тархалт, хугацаа хэтрэлтийн бүтэц, хэтэрсэн ба хэтрээгүй бүлгийн ялгаа, мөн нас, хүйсээр ангилсан зургаан сегментийн эрсдэлийг дэс дараалан дэлгэрэнгүй авч үзнэ.")

    _H1(doc, "1. Онооны тархалт ба ерөнхий шинж")
    _BODY(doc, f"Нийт оноо дунджаар {S['total_score_mean']:.1f} (медиан {S['total_score_med']:.0f}) бөгөөд бодит тархалт [{S['total_score_min']:.0f}–{S['total_score_max']:.0f}] хооронд байна. Ихэнх харилцагч дунд түвшний оноонд төвлөрч, онолын дээд хязгаар (600)-д хүрсэн харилцагч цөөн байна. Нийт оноог бүрдүүлэгч хоёр хэсгээс санхүүгийн оноо дунджаар {S['fin_score_mean']:.1f} ([{S['fin_score_min']:.0f}–{S['fin_score_max']:.0f}]), сэтгэлзүйн оноо дунджаар {S['psy_score_mean']:.1f} ([{S['psy_score_min']:.0f}–{S['psy_score_max']:.0f}]) байгаа нь нийт оноонд санхүүгийн оноо илүү жинтэй нөлөөлж байгааг харуулна. Тархалтын хэлбэр хэвийн (normal)-д ойр, дунд хэсэгтээ хамгийн нягт байна (Зураг 1).")
    _IMG(doc, C["dist"]); _CAP(doc, "Зураг 1. Нийт онооны тархалт (харилцагчийн түвшин)")
    _BODY(doc, f"Хүн ам зүйн бүтцийн хувьд харилцагчдын дундаж нас {S['age_mean']:.1f}, насны бүлгээр 17–29 насныхан нийт харилцагчдын {b17n/S['n_cust']*100:.1f}%, 30–45 насныхан {b30n/S['n_cust']*100:.1f}%, 46-аас дээш насныхан {b46n/S['n_cust']*100:.1f}%-ийг бүрдүүлж байгаа нь багц голчлон залуу хэрэглэгчдэд төвлөрснийг харуулна. Хүйсийн харьцаанд эмэгтэйчүүд {fem:.1f}%, эрэгтэйчүүд {S['male_all']:.1f}% байна. Зан төлөвийн үзүүлэлтээр iOS төхөөрөмж ашиглагч {S['ios_all']:.1f}%, биометр нэвтрэлт идэвхжүүлсэн {S['bio_all']:.1f}% байгаа нь цаашид эрсдэлийн ялгаанд хэрэглэгдэх боломжтой дижитал зан төлөвийн дохио юм.")

    _H1(doc, "2. Хугацаа хэтрэлтийн шинжилгээ")
    _BODY(doc, "Зээл болон харилцагчийн түвшний хугацаа хэтрэлтийн нэгдсэн үзүүлэлтийг доорх хүснэгтэд харьцуулан харуулав.")
    _TABLE(doc, [3360, 3000, 3000], ["Үзүүлэлт", "Зээл (данс)", "Харилцагч"], [
        ["Нийт", _m(S['n_loan']), _m(S['n_cust'])],
        ["Хаалттай (C)", f"{_m(S['nC'])} ({S['nC']/S['n_loan']*100:.1f}%)", "—"],
        ["Нээлттэй", f"{_m(S['nOpen'])} ({S['nOpen']/S['n_loan']*100:.1f}%)", "—"],
        ["Идэвхтэй хэтрэлттэй", f"{_m(S['nOact'])} ({S['nOact']/S['n_loan']*100:.1f}%)", f"{_m(S['od_cust'])} ({S['od_cust_pct']:.1f}%)"],
        ["30+ хоног", _m(S['buckets'][5][1]), f"{_m(S['c30'])} ({S['c30_pct']:.2f}%)"],
    ])
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    _BODY(doc, f"Идэвхтэй хэтрэлттэй {_m(S['nOact'])} дансны дундаж хэтрэлт {S['oa_mean']:.1f} хоног, медиан {S['oa_med']:.0f} хоног байгаа нь хуваарилалт баруун тийш сунасан буюу цөөн тооны урт хугацааны хэтрэлт дунджийг өргөж байгааг илтгэнэ. Хэтрэлтийн хугацаагаар бүлэглэвэл идэвхтэй хэтрэлттэй дансны {SHORT_P:.1f}% нь ердөө 5 хүртэлх хоногийн богино хэтрэлттэй, {MID_P:.1f}% нь 6–15 хоног, харин {LONG_P:.1f}% нь 16-аас дээш хоногийн дунд болон урт хугацааны хэтрэлттэй байна. Богино хугацааны хэтрэлт давамгайлж байгаа нь эерэг хэдий ч 16+ хоногийн дөрөвний нэг орчим хувь нь өндөр эрсдэлийн бүс рүү шилжих магадлалтай тул тусгайлан хянах шаардлагатай. Хамгийн их хэтрэлт {S['oa_max']:.0f} хоног хүрсэн байна (Зураг 2).")
    _IMG(doc, C["bucket"]); _CAP(doc, "Зураг 2. Идэвхтэй хэтрэлтийн бүс (данс, хоногоор)")
    _t2 = f"Хэтрэлтийн бүтцийг үнэлэхэд {_assess_short(SHORT_P)}."
    if prev:
        _t2 += f" Идэвхтэй хэтрэлттэй харилцагчийн эзлэх хувь {_trend(S['od_cust_pct'], prev['od_pct'])}."
    _BODY(doc, _t2)
    _BODY(doc, "Эрсдэлийн оноо хэтрэлтийг хэр сайн ялгаж байгааг үнэлэхийн тулд нийт оноо, санхүүгийн оноо, сэтгэлзүйн оноо тус бүрийг 30 оноогоор жигд бүсчилж, бүс бүрийн харилцагчийн тоо, нийт хэтрэлт % (1+ хоног) болон 15+ хоногийн хэтрэлт %-ийг доор харуулав (Зураг 3–5).")
    _IMG(doc, C["band_total"]); _CAP(doc, "Зураг 3. Нийт онооны бүсээр — харилцагчийн тоо, хэтрэлт % (1+) ба 15+ хэтрэлт %")
    _IMG(doc, C["band_fin"]); _CAP(doc, "Зураг 4. Санхүүгийн онооны бүсээр — харилцагчийн тоо, хэтрэлт % (1+) ба 15+ хэтрэлт %")
    _IMG(doc, C["band_psy"]); _CAP(doc, "Зураг 5. Сэтгэлзүйн онооны бүсээр — харилцагчийн тоо, хэтрэлт % (1+) ба 15+ хэтрэлт %")
    _BODY(doc, f"Гурван онооны ялгах чадварыг харьцуулбал нийт онооны хувьд хамгийн доод бүсэд ({tlo_b}) хэтрэлт {tlo:.1f}% байсан бол хамгийн дээд бүсэд ({thi_b}) {thi:.1f}% хүртэл буурч, {tspr:.1f} нэгж хувийн зөрүү гарч байна. Санхүүгийн оноо үүнтэй ойролцоо ({fspr:.1f}пп зөрүү) ялгах чадвартай бол сэтгэлзүйн оноо хамгийн сул ({pspr:.1f}пп) байна. Өөрөөр хэлбэл хэтрэлтийн эрсдэлийг таамаглахад санхүүгийн болон нийт оноо илүү мэдээлэл өгч байгаа бол сэтгэлзүйн оноо дангаараа сул ялгаатай байна. 15+ хоногийн хэтрэлтийн шугам (ягаан тасархай) нь нийт хэтрэлтийн шугамаас доогуур боловч ижил чиглэлд хөдөлж байгаа нь өндөр эрсдэлтэй хэтрэлт ч мөн оноо буурахад нэмэгддэгийг харуулна. Нийт оноо болон идэвхтэй хэтрэлтийн ерөнхий корреляц {S['corr']:+.3f} буюу сул сөрөг хэвээр байгаа нь оноог дангаар нь биш бусад үзүүлэлттэй хослуулан ашиглах шаардлагатайг дахин нотолж байна.")

    _BODY(doc, f"Хэтрэлтийн босгыг өндөрсгөж үзэхэд +15 хоногийн хэтрэлттэй дансны хувьд нийт онооны корреляц {S['corr15']:+.3f} (n={S['corr15_n']:,}), +30 хоногийнх {S['corr30']:+.3f} (n={S['corr30_n']:,}) байгаа нь өндөр эрсдэлтэй хэтрэлтийн хувьд ч оноотой хамаарал сул хэвээр байгааг харуулна.")
    _BODY(doc, f"Энэ үеийн хувьд хэтрэлтийг хамгийн сайн ялгаж буй үзүүлэлт нь {_best_discriminator(tspr, fspr, pspr)} байгаа бол нийт оноо болон хэтрэлтийн ерөнхий хамаарал {_assess_corr(S['corr'])} ({S['corr']:+.3f}) түвшинд байна.")

    _H1(doc, "3. Хэтэрсэн ба хэтрээгүй бүлгийн харьцуулсан шинжилгээ")
    _BODY(doc, "Хугацаа хэтрэлттэй (идэвхтэй хэтрэлт > 0) болон хэтрэлтгүй харилцагчдыг үзүүлэлт бүрээр харьцуулахад санхүүгийн чадавх болон зан төлөвийн хэд хэдэн хувьсагчид тууштай ялгаа ажиглагдаж байна.")
    cmpn = {"total_score": "Нийт оноо", "fin_score": "Санхүүгийн оноо", "psy_score": "Сэтгэлзүйн оноо", "age": "Нас",
            "slry_last_amt": "Сүүлийн цалин (₮)", "slry_last_avg_6m": "6 сарын дундаж цалин (₮)",
            "zms_active_ln_cnt": "ЗМС идэвхтэй зээл", "zms_monthly_payment": "ЗМС сарын төлбөр (₮)"}
    rows = []
    for c, nm in cmpn.items():
        if c in S["cmp"]:
            o, n = S["cmp"][c]
            money = ("цалин" in nm or "төлбөр" in nm)
            f = (lambda v: f"{v:,.0f}") if money else (lambda v: f"{v:.1f}")
            rows.append([nm, f(o), f(n)])
    _TABLE(doc, [3960, 2700, 2700], ["Үзүүлэлт (дундаж)", "Хэтэрсэн", "Хэтрээгүй"], rows)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    _BODY(doc, "Хэтрэлттэй бүлгийн дундаж нийт оноо, санхүүгийн оноо болон цалингийн бүх үзүүлэлт хэтрэлтгүй бүлгийнхээс тогтмол бага байгаа нь орлогын түвшин болон санхүүгийн тогтвортой байдал хэтрэлтийн эрсдэлийг бууруулдгийг харуулна. Ялгаа том биш ч чиглэл нь нийцтэй байгаа нь эдгээр хүчин зүйлсийн нийлбэр нөлөөг илэрхийлнэ (Зураг 6).")
    _IMG(doc, C["od_score"]); _CAP(doc, "Зураг 6. Дундаж оноо: хэтрээгүй / хэтэрсэн (1+) / 15+ хэтэрсэн бүлэг")
    if bio and cont and male:
        _BODY(doc, f"Зан төлөвийн дохионуудаас хамгийн тод ялгаа нь биометр нэвтрэлтэд ажиглагдаж байна: хэтрэлтгүй бүлгийн {bio[1]:.1f}% биометр идэвхжүүлсэн байхад хэтрэлттэй бүлэгт {bio[0]:.1f}% буюу мэдэгдэхүйц бага байна. Үүнчлэн сүүлийн 3 сарын тасралтгүй цалинтай харилцагч хэтрэлтгүй бүлэгт {cont[1]:.1f}%, хэтрэлттэй бүлэгт {cont[0]:.1f}% байгаа нь орлогын тогтвортой байдлын ач холбогдлыг баталж байна. Жендэрийн хувьд хэтрэлттэй бүлэгт эрэгтэйчүүд {male[0]:.1f}% эзэлж байгаа нь хэтрэлтгүй бүлэг дэх {male[1]:.1f}%-аас өндөр буюу эрэгтэй харилцагчид дунджаар илүү эрсдэлтэй болохыг тодорхой харуулна.")

    _H1(doc, "4. Нас × хүйсний 6 сегментийн шинжилгээ")
    _BODY(doc, "Эрсдэлийн дүр зургийг илүү нарийвчлахын тулд харилцагчдыг 17–29, 30–45, 46+ гэсэн насны гурван бүлэг болон хүйсээр хослуулан зургаан сегментэд хуваав. Доорх хүснэгт болон Зураг 7-д сегмент бүрийн хэмжээ, хэтрэлтийн хувь, дундаж оноо, цалинг харьцуулав.")
    srows = [[s["name"], _m(s["n"]), f"{s['pct']:.1f}%", f"{s['od']:.1f}%", f"{s['tot']:.1f}", f"{s['sal']:.2f}"] for s in S["segs"]]
    _TABLE(doc, [2400, 1320, 1320, 1560, 1380, 1380], ["Сегмент", "Тоо", "Хувь", "Хэтрэлт %", "Нийт оноо", "Цалин (сая₮)"], srows)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    _IMG(doc, C["seg_od"]); _CAP(doc, "Зураг 7. Сегмент бүрийн хэтрэлтийн бүтэц (хэтрээгүй / 1+ / 15+)")
    _BODY(doc, f"Сегментийн шинжилгээгээр хамгийн өндөр эрсдэлтэй бүлэг нь {hi['name']} бөгөөд хэтрэлтийн хувь {hi['od']:.1f}%, 30+ хоногийн хэтрэлт {hi['p30']:.2f}% хүрч байгаа нь нийт дунджаас ({S['od_cust_pct']:.1f}% ба {S['c30_pct']:.2f}%) хамаагүй өндөр байна. Эсрэгээрээ {lo['name']} хамгийн бага эрсдэлтэй ({lo['od']:.1f}%) сегмент юм. Анхаарал татах тууштай хэв маяг нь бүх насны бүлэгт эрэгтэйчүүд эмэгтэйчүүдээс өндөр хэтрэлттэй байгаа явдал (нийтдээ эрэгтэй {em[1]:.1f}% vs эмэгтэй {ef[1]:.1f}%). Насны бүлгээр авч үзвэл 46+ насныхан хамгийн өндөр хэтрэлттэй ({b46o:.1f}%) бол 17–29 ({b17o:.1f}%) болон 30–45 ({b30o:.1f}%) бүлэг ойролцоо байна.")
    _IMG(doc, C["seg_size"]); _CAP(doc, "Зураг 8. Сегментийн хэмжээ (харилцагчийн тоо)")
    _BODY(doc, f"Сегментийн хэмжээний хувьд 17–29 насныхан нийт багцын дийлэнхийг (≈{b17n/S['n_cust']*100:.0f}%) бүрдүүлдэг тул эдгээрийн эрсдэл бага ч нийт хэтрэлтийн дийлэнх тоо энэ бүлгээс гарч байна. Харин {hi['name']} цөөн (≈{_m(seg46m['n'])} харилцагч) хэдий ч хэтрэлтийн хувь өндөр тул нэгж харилцагчид ногдох эрсдэл хамгийн их сегмент юм. Профайлын хувьд 30–45 болон 46+ насны эрэгтэйчүүд санхүүгийн оноо болон цалингаар харьцангуй өндөр атал хэтрэлт өндөр хэвээр байгаа нь зөвхөн санхүүгийн чадавх биш, зээлийн зан төлөв болон орлогын тогтвортой байдал зэрэг бусад хүчин зүйл эрсдэлд нөлөөлж байгааг илэрхийлнэ.")

    _H1(doc, "Дүгнэлт")
    _BODY(doc, f"PsyFint кредит скорингийн загвараар үнэлэгдсэн зээлийн багцын чанар нийтдээ тогтвортой түвшинд байна. Харилцагчдын {S['od_cust_pct']:.1f}% идэвхтэй хэтрэлттэй, 30+ хоногийн өндөр эрсдэлтэй хэсэг {S['c30_pct']:.2f}% буюу харьцангуй бага байгаа нь эерэг үзүүлэлт юм. Идэвхтэй хэтрэлтийн дийлэнх ({SHORT_P:.1f}%) нь 5 хүртэлх хоногийн богино хэтрэлт байгаа нь эргэн төлөлтийн зан төлөв нийтдээ хүлээн зөвшөөрөгдөх түвшинд байгааг харуулна.")
    if prev:
        _BODY(doc, f"Өмнөх үетэй ({prev['name']}) харьцуулсан динамик: идэвхтэй хэтрэлт {_trend(S['od_cust_pct'], prev['od_pct'])}; дундаж нийт оноо {_trend(S['total_score_mean'], prev['score'], unit=' нэгж')}.")
    _BODY(doc, "Анхаарал шаардсан гол дүгнэлтүүд:")
    _BUL(doc, f"{hi['name']} хамгийн өндөр эрсдэлтэй сегмент (хэтрэлт {hi['od']:.1f}%, 30+ {hi['p30']:.2f}%) — тусгайлан хяналт шаардлагатай.")
    _BUL(doc, f"Бүх насны бүлэгт эрэгтэй харилцагчид эмэгтэйчүүдээс тогтмол өндөр хэтрэлттэй ({em[1]:.1f}% vs {ef[1]:.1f}%).")
    _BUL(doc, f"Санхүүгийн болон нийт оноо хэтрэлтийг {fspr:.1f}–{tspr:.1f}пп зөрүүгээр ялгаж байгаа бол сэтгэлзүйн оноо сул ({pspr:.1f}пп) — онооны жинг дахин үнэлэх боломжтой.")
    _BUL(doc, "Биометр нэвтрэлт болон тасралтгүй цалин нь хэтрэлтгүй бүлэгт илүү түгээмэл — эрт сэрэмжлүүлэх дохио болгон ашиглаж болно.")
    _BUL(doc, f"16-аас дээш хоногийн хэтрэлттэй дансууд (нийт {LONG} данс) өндөр эрсдэлийн бүс — тусгай хяналтад авах нь зүйтэй.")
    _BODY(doc, "Цаашид загварын ялгах чадварыг сайжруулахын тулд нас, хүйсээр ялгаатай эрсдэлийн загвар (segment-specific scoring), зан төлөвийн хувьсагчдын жинг нэмэгдүүлэх, мөн санхүүгийн болон сэтгэлзүйн скорингийн жинг оновчтой тохируулах боломжийг авч үзэхийг зөвлөж байна.")

    _add_corr_appendix(doc, S)
    _add_group_profile(doc, S)
    _add_six_segments(doc, S)

    out_path = outdir / f"loan_report_{dt.date.today().strftime('%Y-%m-%d')}.docx"
    doc.save(str(out_path))
    return out_path


def main():
    df, src = load_final_data()
    S = compute(df)
    report = build_report(S, src, OUTPUT_DIR)
    print(f"Тайлан амжилттай хадгалагдлаа: {report}")
    export_for_dashboard(df)


if __name__ == "__main__":
    main()
